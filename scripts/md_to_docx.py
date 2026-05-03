"""Convert the Linux offline discovery enhancement proposal from Markdown to DOCX."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

INPUT = Path(r"c:\Users\Swati Devgan\PM-Skills\output\linux-offline-discovery-enhancement-proposal.md")
OUTPUT = INPUT.with_suffix(".docx")

md = INPUT.read_text(encoding="utf-8")
doc = Document()

# -- styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0, 0, 0)

def add_table_from_rows(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val.strip()
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

lines = md.split("\n")
i = 0
while i < len(lines):
    line = lines[i]

    # Skip horizontal rules
    if re.match(r"^-{3,}$", line.strip()):
        i += 1
        continue

    # Headings
    m = re.match(r"^(#{1,3})\s+(.*)", line)
    if m:
        level = len(m.group(1))
        text = m.group(2).replace("**", "").strip()
        doc.add_heading(text, level=level)
        i += 1
        continue

    # Table detection
    if "|" in line and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
        # Parse header
        headers = [c.strip().replace("**", "") for c in line.strip().strip("|").split("|")]
        i += 2  # skip header + separator
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip().replace("**", "") for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
            i += 1
        add_table_from_rows(headers, rows)
        doc.add_paragraph("")  # spacer
        continue

    # Numbered list items
    m = re.match(r"^(\d+)\.\s+\*\*(.*?)\*\*(.*)", line)
    if m:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(m.group(2))
        run.bold = True
        run.font.size = Pt(11)
        rest = m.group(3).strip()
        if rest:
            run2 = p.add_run(" " + rest)
            run2.font.size = Pt(11)
        i += 1
        continue

    # Bold paragraph lines (like "Phase 1 - Crawl (P0):")
    m = re.match(r"^\*\*(.*?)\*\*(.*)$", line)
    if m and not line.strip().startswith("|"):
        p = doc.add_paragraph()
        run = p.add_run(m.group(1))
        run.bold = True
        run.font.size = Pt(11)
        rest = m.group(2).strip()
        if rest:
            run2 = p.add_run(" " + rest)
            run2.font.size = Pt(11)
        i += 1
        continue

    # Regular paragraph (skip empty)
    text = line.strip()
    if text:
        # Clean markdown formatting
        clean = text.replace("**", "")
        # Handle inline code
        clean = re.sub(r"`([^`]+)`", r"\1", clean)
        p = doc.add_paragraph(clean)
        p.style = doc.styles["Normal"]
        i += 1
        continue

    i += 1

doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
